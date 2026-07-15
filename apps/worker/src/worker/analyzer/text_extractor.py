from __future__ import annotations

import os
import subprocess
import tempfile
import io
from loguru import logger
from worker.resource_limits import positive_int_env


OCR_MAX_PAGES = positive_int_env("WORKER_OCR_MAX_PAGES", 12)
PDF_TEXT_MAX_PAGES = positive_int_env("WORKER_PDF_TEXT_MAX_PAGES", 100)
EXTRACTED_TEXT_MAX_CHARS = positive_int_env("WORKER_EXTRACTED_TEXT_MAX_CHARS", 120_000)


def _bounded_text(text: str) -> str:
    if len(text) <= EXTRACTED_TEXT_MAX_CHARS:
        return text
    logger.warning(
        "Extracted text exceeded {} characters and was truncated",
        EXTRACTED_TEXT_MAX_CHARS,
    )
    return text[:EXTRACTED_TEXT_MAX_CHARS]

def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        text_parts = []
        for i, page in enumerate(reader.pages):
            if i >= PDF_TEXT_MAX_PAGES:
                logger.warning(
                    "PDF {} text extraction is capped at {} pages",
                    file_path,
                    PDF_TEXT_MAX_PAGES,
                )
                break
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        text = "\n".join(text_parts)
    except Exception as e:
        logger.warning(f"Failed to extract text from PDF {file_path}: {e}")

    # OCR Fallback if text is empty or very short
    if len(text.strip()) < 150:
        logger.info(f"PDF {file_path} text content is too short ({len(text)} chars). Trying Tesseract OCR...")
        try:
            import fitz  # PyMuPDF
            import pytesseract
            from PIL import Image

            try:
                pytesseract.get_tesseract_version()
            except Exception:
                logger.warning("Tesseract OCR is not installed or not in PATH — skipping OCR.")
                return text

            ocr_text_parts = []
            with fitz.open(file_path) as doc:
                page_count = min(len(doc), OCR_MAX_PAGES)
                if len(doc) > page_count:
                    logger.warning(
                        "PDF {} has {} pages; OCR is capped at {} pages",
                        file_path,
                        len(doc),
                        page_count,
                    )
                for i in range(page_count):
                    page = doc.load_page(i)
                    mat = fitz.Matrix(2.0, 2.0)
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("png")
                    with Image.open(io.BytesIO(img_data)) as img:
                        page_ocr = pytesseract.image_to_string(
                            img,
                            lang="rus+bel+eng",
                        )
                    del img_data, pix, page
                    if page_ocr.strip():
                        ocr_text_parts.append(page_ocr)

            if ocr_text_parts:
                logger.info(f"OCR successfully extracted {len(ocr_text_parts)} pages from PDF {file_path}")
                return _bounded_text("\n".join(ocr_text_parts))
        except Exception as ocr_err:
            logger.warning(f"OCR failed for {file_path}: {ocr_err}")

    return text

def extract_text_from_docx(file_path: str) -> str:
    try:
        import docx
        doc = docx.Document(file_path)
        text_parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                text_parts.append(" | ".join(row_text))
        return "\n".join(text_parts)
    except Exception as e:
        logger.warning(f"Failed to extract text from DOCX {file_path}: {e}")
        return ""

def extract_text_from_doc(file_path: str) -> str:
    try:
        # Try antiword first (handles real binary .doc)
        result = subprocess.run(
            ["antiword", file_path],
            capture_output=True,
            text=True,
            check=True,
            timeout=10,
        )
        return result.stdout
    except Exception as e:
        logger.warning(f"antiword failed to parse {file_path}, falling back to docx parser: {e}")
        # Fallback to python-docx in case it is a renamed docx/xml
        return extract_text_from_docx(file_path)

def extract_text_from_xlsx(file_path: str) -> str:
    try:
        from openpyxl import load_workbook
        wb = load_workbook(file_path, read_only=True, data_only=True)
        text_parts: list[str] = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(
                    str(cell) for cell in row if cell is not None and str(cell).strip()
                )
                if row_text:
                    text_parts.append(row_text)
        wb.close()
        return "\n".join(text_parts)
    except Exception as e:
        logger.warning(f"Failed to extract text from XLSX {file_path}: {e}")
        return ""

def extract_text_from_xls(file_path: str) -> str:
    try:
        import xlrd
        wb = xlrd.open_workbook(file_path)
        text_parts = []
        for sheet in wb.sheets():
            for row_idx in range(sheet.nrows):
                row_vals = sheet.row_values(row_idx)
                row_text = " | ".join(
                    str(val) for val in row_vals if val is not None and str(val).strip()
                )
                if row_text:
                    text_parts.append(row_text)
        return "\n".join(text_parts)
    except Exception as e:
        logger.warning(f"Failed to extract text from XLS {file_path}: {e}")
        return ""

def extract_text_from_archive(file_path: str) -> str:
    _, ext = os.path.splitext(file_path.lower())
    extracted_text_parts = []

    with tempfile.TemporaryDirectory() as temp_dir:
        try:
            if ext == ".zip":
                import zipfile
                with zipfile.ZipFile(file_path, "r") as z:
                    z.extractall(temp_dir)
            elif ext == ".rar":
                import rarfile
                try:
                    with rarfile.RarFile(file_path, "r") as r:
                        r.extractall(temp_dir)
                except Exception as rar_e:
                    logger.warning(f"rarfile failed (unrar might be missing): {rar_e}")
            elif ext == ".7z":
                import py7zr
                with py7zr.SevenZipFile(file_path, "r") as s:
                    s.extractall(temp_dir)
            else:
                return ""

            # Recursively walk through the files
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    sub_file_path = os.path.join(root, file)
                    sub_ext = os.path.splitext(file.lower())[1]
                    if sub_ext in (".pdf", ".docx", ".doc", ".xlsx", ".xls"):
                        sub_text = extract_text_from_file(sub_file_path)
                        if sub_text.strip():
                            extracted_text_parts.append(
                                f"--- Extracted from {file} ---\n{sub_text}"
                            )
        except Exception as e:
            logger.warning(f"Failed to extract archive {file_path}: {e}")

    return "\n\n".join(extracted_text_parts)

def extract_text_from_file(file_path: str) -> str:
    _, ext = os.path.splitext(file_path.lower())
    if ext == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif ext == ".docx":
        text = extract_text_from_docx(file_path)
    elif ext == ".doc":
        text = extract_text_from_doc(file_path)
    elif ext == ".xlsx":
        text = extract_text_from_xlsx(file_path)
    elif ext == ".xls":
        text = extract_text_from_xls(file_path)
    elif ext in (".zip", ".rar", ".7z"):
        text = extract_text_from_archive(file_path)
    else:
        logger.warning(f"Unsupported file extension {ext} for {file_path}")
        text = ""
    return _bounded_text(text)
